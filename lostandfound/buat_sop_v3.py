from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
s = doc.sections[0]
s.top_margin = Inches(1); s.bottom_margin = Inches(1)
s.left_margin = Inches(1.2); s.right_margin = Inches(1.2)

def h1(doc, text, color=(79,70,229)):
    p = doc.add_heading(text, level=1)
    for r in p.runs: r.font.color.rgb = RGBColor(*color); r.font.size = Pt(14)

def h2(doc, text, color=(30,58,138)):
    p = doc.add_heading(text, level=2)
    for r in p.runs: r.font.color.rgb = RGBColor(*color); r.font.size = Pt(12)

def b(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    p.runs[0].font.size = Pt(11)

def note(doc, text, color=(146,64,14)):
    p = doc.add_paragraph(text)
    for r in p.runs: r.font.size=Pt(10); r.font.color.rgb=RGBColor(*color); r.italic=True

# TITLE
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('SOP APLIKASI LOST & FOUND — BXRINK ICE SKATING')
r.bold=True; r.font.size=Pt(20); r.font.color.rgb=RGBColor(79,70,229)
p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('Standar Operasional Prosedur — Versi Terbaru')
r2.italic=True; r2.font.size=Pt(12); r2.font.color.rgb=RGBColor(107,114,128)
doc.add_paragraph()

# A
h1(doc, 'A. LEVEL AKSES PENGGUNA')

h2(doc, '👤 Guest (Tanpa Login)', (100,116,139))
b(doc, 'Lihat daftar Barang Hilang & Barang Ditemukan')
b(doc, 'Lapor Barang Hilang TANPA perlu login')
b(doc, 'Lihat panduan cara klaim barang')
note(doc, 'TIDAK DAPAT: Ajukan klaim, cek status, lihat foto & kontak pelapor')

h2(doc, '🙋 Rink Operator (Login Required)', (37,99,235))
b(doc, 'Semua akses Guest')
b(doc, 'Lapor Barang Ditemukan')
b(doc, 'Ajukan Klaim Pengambilan (langsung selesai tanpa verifikasi admin)')
b(doc, 'Konfirmasi serah terima & upload foto bukti')
b(doc, 'Lihat foto barang dan kontak pelapor')
note(doc, 'TIDAK DAPAT: Panel Admin, Pemusnahan, Semua Laporan, Riwayat')

h2(doc, '👑 Admin / Superadmin (Login Required)', (79,70,229))
b(doc, 'Semua akses Rink Operator')
b(doc, 'Panel Admin: Dashboard, Semua Laporan, Klaim & Pengambilan')
b(doc, 'Pemusnahan & Donasi barang')
b(doc, 'Export data CSV & Sync Google Sheets otomatis')
b(doc, 'Kelola Pengguna & Riwayat Pengambilan')

h2(doc, '🔐 Cara Login', (79,70,229))
b(doc, 'Klik tombol 👤 Masuk di navbar')
b(doc, 'Masukkan Username dan Password')
b(doc, 'Selesaikan reCAPTCHA')
b(doc, 'Klik Masuk')
note(doc, 'Username: superadmin (Admin) · rinkoperator (Rink Operator) | Akses Admin: Ctrl+Shift+A')
doc.add_paragraph()

# B
h1(doc, 'B. LAPORAN BARANG HILANG (Semua Pengguna — Tanpa Login)')
b(doc, 'B1. Buka halaman 🚨 Barang Hilang di navbar')
b(doc, 'B2. Klik tombol + Laporkan Hilang (tanpa perlu login)')
b(doc, 'B3. Isi: nama barang, kategori, lokasi, deskripsi, foto (opsional), kontak')
b(doc, 'B4. Selesaikan reCAPTCHA')
b(doc, 'B5. Klik Kirim Laporan')
note(doc, 'Sistem otomatis mencari kandidat barang ditemukan yang cocok berdasarkan deskripsi & foto')
doc.add_paragraph()

# C
h1(doc, 'C. LAPORAN BARANG DITEMUKAN (Login Required)')
b(doc, 'C1. Login sebagai Rink Operator atau Admin')
b(doc, 'C2. Buka halaman ✅ Barang Ditemukan')
b(doc, 'C3. Klik + Laporkan Temuan')
b(doc, 'C4. Isi: nama barang, kategori, lokasi ditemukan, deskripsi, foto, nama penemu, kontak')
b(doc, 'C5. Klik Kirim Laporan')
note(doc, 'Sistem otomatis mencari kandidat barang hilang yang cocok & membandingkan foto')
doc.add_paragraph()

# D
h1(doc, 'D. PROSES KLAIM PENGAMBILAN (Login Required — Langsung Selesai)')
p = doc.add_paragraph()
r = p.add_run('⚡ TIDAK PERLU VERIFIKASI ADMIN — Klaim langsung selesai saat disubmit!')
r.bold=True; r.font.color.rgb=RGBColor(220,38,38); r.font.size=Pt(11)

h2(doc, 'D1 — Step 1: Pilih Barang', (30,58,138))
b(doc, 'Buka menu 🤝 Pengambilan → klik tab Ajukan Klaim')
b(doc, 'Cari dan pilih barang yang ingin diklaim')
b(doc, 'Klik Lanjut')

h2(doc, 'D2 — Step 2: Bukti Kepemilikan', (30,58,138))
b(doc, 'Isi nama pengklaim')
b(doc, 'Isi nomor kontak (WA)')
b(doc, 'Isi nama petugas')
b(doc, 'Deskripsikan bukti kepemilikan secara detail')
b(doc, 'Upload foto diri pengambil sebagai bukti')
b(doc, 'Klik Lanjut')

h2(doc, 'D3 — Step 3: Jadwal Pengambilan', (30,58,138))
b(doc, 'Isi tanggal pengambilan barang')
b(doc, 'Pilih waktu (08:00 - 22:00)')
b(doc, 'Isi catatan tambahan jika diperlukan')
b(doc, 'Klik Kirim Klaim')
note(doc, 'Status klaim: Completed | Barang: Resolved | Barang hilang dari daftar otomatis')
doc.add_paragraph()

# E
h1(doc, 'E. KONFIRMASI SERAH TERIMA (Admin atau Rink Operator)')
b(doc, 'E1. Buka menu 🤝 Pengambilan → tab Cek Status')
b(doc, 'E2. Cari klaim yang sudah completed')
b(doc, 'E3. Klik Upload Bukti Serah Terima')
b(doc, 'E4. Upload foto serah terima')
b(doc, 'E5. Isi catatan serah terima (opsional)')
b(doc, 'E6. Klik Konfirmasi Selesai')
note(doc, 'Bukti foto tersimpan di database dan tercatat di Riwayat Pengambilan (Admin)')
doc.add_paragraph()

# F
h1(doc, 'F. PEMUSNAHAN / DONASI (Admin Only — Barang Tidak Diklaim > 30 Hari)', (220,38,38))
b(doc, 'F1. Buka menu 🗑️ Pemusnahan')
b(doc, 'F2. Cek barang eligible (open, tidak ada klaim, lebih dari 30 hari)')
b(doc, 'F3. Pilih barang → klik Proses Pemusnahan / Donasi')
b(doc, 'F4. Pilih: 🔥 Dimusnahkan atau 💙 Didonasikan')
b(doc, 'F5. Isi formulir: tanggal, petugas, alasan, saksi')
b(doc, 'F6. Upload foto bukti tindakan (WAJIB)')
b(doc, 'F7. Centang pernyataan konfirmasi → klik Proses Tindakan')
note(doc, 'Status: Dimusnahkan/Didonasikan | Tercatat di Riwayat Pemusnahan')
doc.add_paragraph()

# G
h1(doc, 'G. FITUR PANEL ADMIN (Login Superadmin)')

h2(doc, '📊 Dashboard', (79,70,229))
b(doc, 'Statistik laporan per bulan (grafik bar)')
b(doc, 'Grafik status laporan & kategori barang')
b(doc, 'Lokasi dengan kejadian terbanyak')
b(doc, 'Total klaim selesai')

h2(doc, '📋 Semua Laporan', (79,70,229))
b(doc, 'Timeline semua barang hilang & ditemukan')
b(doc, 'Filter berdasarkan status, kategori')
b(doc, 'Konfirmasi auto-match antara barang hilang & ditemukan')

h2(doc, '⬇️ Export & Sync Data', (79,70,229))
b(doc, 'Export CSV: Barang Hilang, Barang Ditemukan, Data Klaim')
b(doc, 'Sync otomatis ke Google Sheets setiap ada perubahan data')

h2(doc, '👥 Kelola Pengguna', (79,70,229))
b(doc, 'Lihat daftar pengguna terdaftar')
b(doc, 'Ubah role: Rink Operator ↔ Admin')
b(doc, 'Riwayat Pengambilan (foto bukti serah terima)')
b(doc, 'Hapus laporan barang')
doc.add_paragraph()

# H
h1(doc, 'H. INFORMASI PENTING')

h2(doc, '📍 Lokasi Area BXRink', (8,145,178))
b(doc, 'Ice Rink Arena')
b(doc, 'Member Room Figure · Member Room Hockey')
b(doc, 'Multifunction Room / Off Ice Room')
b(doc, 'Waiting Area · Prayer Room · Toilet · Other')

h2(doc, '📁 Kategori Barang', (8,145,178))
b(doc, 'Blade Guard · Botol Minum · Dompet / Tas')
b(doc, 'Elektronik · Hockey Gear · Kunci · Pakaian · Lainnya')

h2(doc, '📞 Kontak & Akses', (22,163,74))
b(doc, 'WhatsApp: +62 812-3456-7890')
b(doc, 'Jam: Senin – Minggu, 10.00 – 17.00 WIB')
b(doc, 'URL: https://stirring-yeot-a170188.netlify.app')
b(doc, 'Data: Firebase Firestore + Google Sheets sync')

out = r'C:\Users\Lenovo\OneDrive\Documents\lostandfound\SOP_LostFound_v3.docx'
doc.save(out)
print('SUCCESS:', out)
