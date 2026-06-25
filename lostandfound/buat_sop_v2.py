from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
section = doc.sections[0]
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin   = Inches(1.2)
section.right_margin  = Inches(1.2)

def h1(doc, text, color=(79,70,229)):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = RGBColor(*color)
        run.font.size = Pt(14)
    return p

def h2(doc, text, color=(30,58,138)):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = RGBColor(*color)
        run.font.size = Pt(12)
    return p

def bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    p.runs[0].font.size = Pt(11)
    return p

def info(doc, text, color=(107,114,128)):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(*color)
        run.italic = True
    return p

# TITLE
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('SOP PENGGUNAAN APLIKASI LOST & FOUND')
run.bold = True; run.font.size = Pt(22); run.font.color.rgb = RGBColor(79,70,229)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('BXRink Ice Skating')
run.bold = True; run.font.size = Pt(16); run.font.color.rgb = RGBColor(67,56,202)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Standar Operasional Prosedur Penggunaan Sistem')
run.italic = True; run.font.size = Pt(11); run.font.color.rgb = RGBColor(107,114,128)
doc.add_paragraph()

# A
h1(doc, 'A. AKSES & LEVEL PENGGUNA')

h2(doc, '👤 Guest (Tanpa Login)', (100,116,139))
bullet(doc, 'Lihat daftar Barang Hilang & Barang Ditemukan')
bullet(doc, 'Lapor Barang Hilang (tanpa login)')
bullet(doc, 'Lihat panduan cara klaim barang')
info(doc, 'TIDAK DAPAT: Ajukan klaim, cek status klaim, lihat foto barang, lihat kontak pelapor')

h2(doc, '🙋 User / Rink Operator (Login Required)', (37,99,235))
bullet(doc, 'Semua akses Guest')
bullet(doc, 'Lapor Barang Ditemukan')
bullet(doc, 'Ajukan Klaim & Cek Status Klaim')
bullet(doc, 'Lihat foto barang dan kontak pelapor')
info(doc, 'TIDAK DAPAT: Panel Admin, Riwayat Pengambilan, Semua Laporan, Pemusnahan')

h2(doc, '👑 Admin / Superadmin (Login Required)', (79,70,229))
bullet(doc, 'Semua akses User')
bullet(doc, 'Panel Admin: kelola semua laporan, klaim, dashboard')
bullet(doc, 'Approve / Reject klaim & jadwalkan pengambilan')
bullet(doc, 'Pemusnahan & Donasi barang')
bullet(doc, 'Export data CSV & Kelola pengguna')
bullet(doc, 'Riwayat Pengambilan & Semua Laporan')

h2(doc, '🔐 Cara Login', (79,70,229))
bullet(doc, 'Klik tombol 👤 Masuk di pojok kanan atas navbar')
bullet(doc, 'Masukkan Username dan Password')
bullet(doc, 'Selesaikan verifikasi reCAPTCHA')
bullet(doc, 'Klik Masuk')
info(doc, 'Username: superadmin (Admin) · rinkoperator (User) — Password diberikan pengelola sistem')
doc.add_paragraph()

# B
h1(doc, 'B. LAPORAN BARANG HILANG (Semua Pengguna — Tanpa Login)')
h2(doc, 'Langkah-langkah:', (30,58,138))
bullet(doc, 'B1. Buka halaman Barang Hilang (klik 🚨 di navbar)')
bullet(doc, 'B2. Klik tombol + Laporkan Hilang')
bullet(doc, 'B3. Isi formulir: nama barang, kategori, lokasi, deskripsi lengkap, foto, kontak')
bullet(doc, 'B4. Selesaikan reCAPTCHA')
bullet(doc, 'B5. Klik Kirim Laporan')
info(doc, 'Laporan tersimpan otomatis & sistem akan mencari kandidat barang ditemukan yang cocok')
doc.add_paragraph()

# C
h1(doc, 'C. LAPORAN BARANG DITEMUKAN (Login Required)')
h2(doc, 'Langkah-langkah:', (30,58,138))
bullet(doc, 'C1. Login menggunakan akun User atau Admin')
bullet(doc, 'C2. Buka halaman Barang Ditemukan (klik ✅ di navbar)')
bullet(doc, 'C3. Klik tombol + Laporkan Temuan')
bullet(doc, 'C4. Isi formulir: nama barang, kategori, lokasi ditemukan, deskripsi, foto, nama penemu, kontak')
bullet(doc, 'C5. Klik Kirim Laporan')
info(doc, 'Sistem otomatis mencari kandidat barang hilang yang cocok berdasarkan deskripsi & foto')
doc.add_paragraph()

# D
h1(doc, 'D. PROSES KLAIM PENGAMBILAN (Login Required)')

h2(doc, 'D1 - Step 1: Pilih Barang', (30,58,138))
bullet(doc, 'Buka menu 🤝 Pengambilan')
bullet(doc, 'Klik tab Ajukan Klaim')
bullet(doc, 'Cari dan pilih barang yang ingin diklaim')
bullet(doc, 'Klik Lanjut')

h2(doc, 'D2 - Step 2: Bukti Kepemilikan', (30,58,138))
bullet(doc, 'Isi nama pengklaim')
bullet(doc, 'Isi nomor kontak (WA)')
bullet(doc, 'Isi nama petugas')
bullet(doc, 'Deskripsikan bukti kepemilikan secara detail')
bullet(doc, 'Bubuhkan tanda tangan digital')
bullet(doc, 'Klik Lanjut')

h2(doc, 'D3 - Step 3: Jadwal Pengambilan', (30,58,138))
bullet(doc, 'Pilih tanggal dan waktu yang diinginkan')
bullet(doc, 'Isi catatan tambahan jika diperlukan')
bullet(doc, 'Klik Kirim Klaim')
info(doc, 'Status klaim menjadi: Menunggu Verifikasi Admin')

h2(doc, 'D4 - Admin Review & Keputusan', (79,70,229))
bullet(doc, 'Admin buka Panel Admin → tab Klaim & Pengambilan')
bullet(doc, 'Review klaim: nama pengklaim, bukti kepemilikan, tanda tangan digital')
bullet(doc, 'DISETUJUI → Isi tanggal, waktu, lokasi, petugas → kode verifikasi otomatis dibuat')
bullet(doc, 'DITOLAK → Isi alasan penolakan → status klaim berubah menjadi Ditolak')
doc.add_paragraph()

# E
h1(doc, 'E. SERAH TERIMA BARANG')
bullet(doc, 'E1. Pengklaim datang ke BXRink sesuai jadwal yang dikonfirmasi')
bullet(doc, 'E2. Tunjukkan kode verifikasi kepada petugas')
bullet(doc, 'E3. Petugas mencocokkan kode dengan sistem')
bullet(doc, 'E4. Admin klik Konfirmasi Sudah Diambil di panel')
bullet(doc, 'E5. Upload foto serah terima sebagai bukti')
bullet(doc, 'E6. Status laporan menjadi Selesai — tersimpan di Riwayat Pengambilan')
doc.add_paragraph()

# F
h1(doc, 'F. PEMUSNAHAN / DONASI BARANG (Admin Only)', (220,38,38))
p = doc.add_paragraph('Kondisi: Barang tidak diklaim selama lebih dari 30 hari')
for run in p.runs:
    run.bold = True; run.italic = True; run.font.color.rgb = RGBColor(220,38,38)

bullet(doc, 'F1. Buka menu Pemusnahan')
bullet(doc, 'F2. Cek daftar barang eligible (open, tidak ada klaim aktif, lebih dari 30 hari)')
bullet(doc, 'F3. Pilih barang → klik Proses Pemusnahan / Donasi')
bullet(doc, 'F4. Pilih jenis tindakan: Dimusnahkan atau Didonasikan')
bullet(doc, 'F5. Isi formulir: tanggal, petugas pelaksana, alasan, saksi')
bullet(doc, 'F6. Upload foto bukti tindakan (wajib)')
bullet(doc, 'F7. Centang pernyataan konfirmasi → klik Proses Tindakan')
info(doc, 'Status barang menjadi Dimusnahkan atau Didonasikan, tercatat di Riwayat Pemusnahan')
doc.add_paragraph()

# G
h1(doc, 'G. FITUR PANEL ADMIN')

h2(doc, '📊 Dashboard', (79,70,229))
bullet(doc, 'Statistik laporan per bulan (grafik bar)')
bullet(doc, 'Grafik status laporan & kategori barang')
bullet(doc, 'Lokasi dengan kejadian terbanyak')
bullet(doc, 'Total klaim: pending, disetujui, selesai')

h2(doc, '📋 Semua Laporan', (79,70,229))
bullet(doc, 'Timeline semua barang hilang & ditemukan')
bullet(doc, 'Filter berdasarkan status, kategori, tanggal')
bullet(doc, 'Konfirmasi auto-match antara barang hilang & ditemukan')

h2(doc, '⬇️ Export Data (CSV)', (79,70,229))
bullet(doc, 'Export CSV Barang Hilang')
bullet(doc, 'Export CSV Barang Ditemukan')
bullet(doc, 'Export CSV Data Klaim')
info(doc, 'File CSV dapat dibuka di Microsoft Excel atau Google Sheets')

h2(doc, '👥 Kelola Pengguna', (79,70,229))
bullet(doc, 'Lihat daftar semua pengguna terdaftar')
bullet(doc, 'Ubah role: User menjadi Admin atau sebaliknya')
bullet(doc, 'Hapus laporan barang')
doc.add_paragraph()

# H
h1(doc, 'H. INFORMASI PENTING')

h2(doc, '📍 Lokasi Area BXRink', (8,145,178))
bullet(doc, 'Ice Rink Arena')
bullet(doc, 'Member Room Figure')
bullet(doc, 'Member Room Hockey')
bullet(doc, 'Multifunction Room / Off Ice Room')
bullet(doc, 'Waiting Area')
bullet(doc, 'Prayer Room')
bullet(doc, 'Toilet')
bullet(doc, 'Other')

h2(doc, '📁 Kategori Barang', (8,145,178))
bullet(doc, 'Blade Guard · Botol Minum · Dompet / Tas')
bullet(doc, 'Elektronik · Hockey Gear · Kunci · Pakaian · Lainnya')

h2(doc, '📞 Hubungi Kami', (22,163,74))
bullet(doc, 'Telepon / WhatsApp: +62 812-3456-7890')
bullet(doc, 'Jam operasional: Senin – Minggu, 10.00 – 17.00 WIB')

h2(doc, '🌐 Akses Aplikasi', (79,70,229))
bullet(doc, 'URL: https://stirring-yeot-a170188.netlify.app')
bullet(doc, 'Data tersimpan di Firebase Firestore (cloud)')
bullet(doc, 'reCAPTCHA aktif pada form laporan barang')

out = r'C:\Users\Lenovo\OneDrive\Documents\lostandfound\SOP_LostFound_v2.docx'
doc.save(out)
print('SUCCESS:', out)
