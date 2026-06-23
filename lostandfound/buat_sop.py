from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Page margins
section = doc.sections[0]
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin   = Inches(1.2)
section.right_margin  = Inches(1.2)

# Style helpers
def add_title(doc, text, color=(79,70,229)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(*color)
    return p

def add_h1(doc, text, color=(79,70,229)):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = RGBColor(*color)
        run.font.size = Pt(14)
    return p

def add_h2(doc, text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = RGBColor(30,58,138)
        run.font.size = Pt(12)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    p.runs[0].font.size = Pt(11)
    return p

def add_space(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)

# ── TITLE ──
add_title(doc, 'SOP ALUR SISTEM LOST & FOUND')
add_title(doc, 'ICE SKATING', (67,56,202))
p = doc.add_paragraph('Standar Operasional Prosedur dari Barang Ditemukan hingga Diselesaikan')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in p.runs:
    run.italic = True
    run.font.color.rgb = RGBColor(107,114,128)
    run.font.size = Pt(11)
doc.add_paragraph()

# ── FASE 1 ──
add_h1(doc, 'FASE 1 — PENEMUAN & PELAPORAN BARANG')
add_h2(doc, 'Langkah 1: Barang Ditemukan')
add_bullet(doc, 'Petugas atau tamu menemukan barang di area ice rink')
add_bullet(doc, 'Barang tidak boleh dibiarkan atau dipindahkan tanpa laporan')
add_h2(doc, 'Langkah 2: Serahkan ke Petugas')
add_bullet(doc, 'Barang diserahkan ke staff atau security terdekat')
add_bullet(doc, 'Petugas mencatat kondisi barang secara fisik')
add_h2(doc, 'Langkah 3: Input Laporan Barang Ditemukan di Aplikasi')
add_bullet(doc, 'Login menggunakan akun yang tersedia')
add_bullet(doc, 'Buka menu Barang Ditemukan → klik + Laporkan Temuan')
add_bullet(doc, 'Isi: nama barang, kategori, lokasi ditemukan, deskripsi lengkap, foto barang, nama penemu, kontak')
add_bullet(doc, 'Klik Kirim Laporan')
add_h2(doc, 'Langkah 4: Sistem Cek Auto-Match')
add_bullet(doc, 'Sistem otomatis membandingkan dengan laporan barang hilang yang ada')
add_bullet(doc, 'Jika ada kandidat cocok → tampil panel "Barang Hilang yang Mungkin Cocok"')
add_bullet(doc, 'Admin dapat konfirmasi kecocokan langsung dari panel tersebut')
add_space(doc)

# ── FASE 2 ──
add_h1(doc, 'FASE 2 — PENCARIAN PEMILIK')
add_h2(doc, 'Langkah 5: Barang Disimpan Petugas')
add_bullet(doc, 'Barang disimpan di tempat yang aman')
add_bullet(doc, 'Tunggu pemilik menghubungi atau mencari melalui aplikasi')
add_h2(doc, 'Langkah 6a: Pemilik Lapor Kehilangan (via Aplikasi)')
add_bullet(doc, 'Pemilik membuka menu Barang Hilang → klik + Laporkan Hilang')
add_bullet(doc, 'Isi: nama barang, kategori, lokasi kejadian, deskripsi, foto (jika ada), nama pelapor, kontak')
add_h2(doc, 'Langkah 6b: Pemilik Mencari Langsung')
add_bullet(doc, 'Buka menu Barang Ditemukan')
add_bullet(doc, 'Ketik nama atau deskripsi barang di kolom pencarian')
add_bullet(doc, 'Sistem akan menampilkan barang yang paling sesuai')
add_space(doc)

# ── FASE 3 ──
add_h1(doc, 'FASE 3 — PROSES KLAIM PENGAMBILAN')
add_h2(doc, 'Langkah 7: Ajukan Klaim — Step 1 (Pilih Barang)')
add_bullet(doc, 'Buka menu Pengambilan → klik Ajukan Klaim')
add_bullet(doc, 'Cari dan pilih barang yang ingin diklaim')
add_bullet(doc, 'Klik Lanjut')
add_h2(doc, 'Langkah 8: Bukti Kepemilikan — Step 2')
add_bullet(doc, 'Isi nama pengklaim')
add_bullet(doc, 'Isi nomor kontak (WA)')
add_bullet(doc, 'Isi nama petugas yang menyerahkan')
add_bullet(doc, 'Deskripsikan bukti kepemilikan secara detail')
add_bullet(doc, 'Bubuhkan tanda tangan digital')
add_bullet(doc, 'Klik Lanjut')
add_h2(doc, 'Langkah 9: Jadwal Pengambilan — Step 3')
add_bullet(doc, 'Pilih tanggal dan waktu yang diinginkan')
add_bullet(doc, 'Isi catatan tambahan jika diperlukan')
add_bullet(doc, 'Klik Kirim Klaim')
add_bullet(doc, 'Status klaim menjadi: ⏳ Menunggu Verifikasi')
add_space(doc)

# ── FASE 4 ──
add_h1(doc, 'FASE 4 — VERIFIKASI ADMIN')
add_h2(doc, 'Langkah 10: Admin Review Klaim')
add_bullet(doc, 'Login sebagai admin (username: superadmin)')
add_bullet(doc, 'Buka menu Admin → tab Klaim & Pengambilan')
add_bullet(doc, 'Review klaim: nama pengklaim, bukti kepemilikan, tanda tangan')
add_h2(doc, 'Langkah 11a: Klaim Disetujui')
add_bullet(doc, 'Admin klik "Setujui & Jadwalkan"')
add_bullet(doc, 'Isi tanggal, waktu, lokasi, dan nama petugas penyerah')
add_bullet(doc, 'Kode verifikasi otomatis dibuat oleh sistem')
add_bullet(doc, 'Status menjadi: 📅 Dijadwalkan')
add_bullet(doc, 'Pengklaim dapat melihat tiket pengambilan di menu Cek Status')
add_h2(doc, 'Langkah 11b: Klaim Ditolak')
add_bullet(doc, 'Admin klik "Tolak" dan isi alasan penolakan')
add_bullet(doc, 'Status menjadi: ❌ Ditolak')
add_bullet(doc, 'Pengklaim mendapat notifikasi alasan penolakan')
add_space(doc)

# ── FASE 5 ──
add_h1(doc, 'FASE 5 — SERAH TERIMA & PENYELESAIAN')
add_h2(doc, 'Langkah 12: Pemilik Datang Mengambil Barang')
add_bullet(doc, 'Pemilik datang sesuai jadwal yang telah dikonfirmasi')
add_bullet(doc, 'Tunjukkan kode verifikasi kepada petugas')
add_bullet(doc, 'Petugas mencocokkan kode verifikasi dengan sistem')
add_h2(doc, 'Langkah 13: Admin Konfirmasi Pengambilan')
add_bullet(doc, 'Buka panel Admin → Klaim & Pengambilan')
add_bullet(doc, 'Klik "Konfirmasi Sudah Diambil"')
add_bullet(doc, 'Upload foto serah terima sebagai bukti')
add_bullet(doc, 'Isi catatan serah terima')
add_bullet(doc, 'Status laporan menjadi: ✅ Selesai')
add_bullet(doc, 'Data tersimpan di Firebase dan tercatat di Riwayat Pengambilan')
add_space(doc)

# ── FASE ALT ──
add_h1(doc, 'FASE ALTERNATIF — TIDAK ADA PEMILIK (PEMUSNAHAN / DONASI)', (220,38,38))
p = doc.add_paragraph('Kondisi: Barang tidak diklaim selama lebih dari 30 hari')
for run in p.runs:
    run.bold = True
    run.italic = True
    run.font.color.rgb = RGBColor(220,38,38)
add_h2(doc, 'Langkah A1: Cek Barang Eligible')
add_bullet(doc, 'Buka menu Pemusnahan')
add_bullet(doc, 'Lihat daftar barang yang memenuhi kriteria (open, tidak ada klaim aktif, lebih dari 30 hari)')
add_h2(doc, 'Langkah A2: Proses Pemusnahan atau Donasi')
add_bullet(doc, 'Pilih barang → klik "Proses Pemusnahan / Donasi"')
add_bullet(doc, 'Pilih jenis tindakan: Dimusnahkan atau Didonasikan')
add_bullet(doc, 'Jika donasi: isi nama dan alamat penerima')
add_bullet(doc, 'Isi tanggal, petugas pelaksana, alasan tindakan, saksi')
add_bullet(doc, 'Upload foto bukti tindakan')
add_bullet(doc, 'Centang pernyataan konfirmasi → klik "Proses Tindakan"')
add_bullet(doc, 'Status barang menjadi: 🔥 Dimusnahkan atau 💙 Didonasikan')
add_bullet(doc, 'Tercatat di Riwayat Pemusnahan')
add_space(doc)

# ── AKUN ──
add_h1(doc, 'INFORMASI AKUN SISTEM')
add_h2(doc, 'Akun Admin')
add_bullet(doc, 'Username: superadmin')
add_bullet(doc, 'Akses: Semua fitur (Semua Laporan, Pemusnahan, Panel Admin, Kelola Pengguna, Riwayat Pengambilan)')
add_h2(doc, 'Akun User / Operator')
add_bullet(doc, 'Username: rinkoperator')
add_bullet(doc, 'Akses: Lapor barang, Ajukan klaim, Cek status klaim')
add_h2(doc, 'Guest (Tanpa Login)')
add_bullet(doc, 'Akses: Lihat daftar barang, Informasi kontak')
add_bullet(doc, 'Tidak dapat: Lapor barang, Ajukan klaim, Lihat foto barang')
add_space(doc)

# ── KONTAK ──
add_h1(doc, 'INFORMASI KONTAK')
add_bullet(doc, 'Hubungi petugas: +62 812-3456-7890 (WhatsApp)')
add_bullet(doc, 'Jam operasional: Senin – Minggu, 10.00 – 17.00 WIB')

# Save
out = r'C:\Users\Lenovo\OneDrive\Documents\lostandfound\SOP_Lost_Found.docx'
doc.save(out)
print('SUCCESS:', out)
