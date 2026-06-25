from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
s = doc.sections[0]
s.page_width  = Inches(8.27)   # A4
s.page_height = Inches(11.69)
s.top_margin = Inches(1); s.bottom_margin = Inches(1)
s.left_margin = Inches(1.2); s.right_margin = Inches(1.2)

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def title_para(doc, text, size=24, color=(79,70,229), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER):
    p = doc.add_paragraph(); p.alignment = align
    r = p.add_run(text); r.bold=bold
    r.font.size=Pt(size); r.font.color.rgb=RGBColor(*color)
    return p

def h1(doc, emoji, text, color=(79,70,229)):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(f'{emoji}  {text}')
    r.bold=True; r.font.size=Pt(15); r.font.color.rgb=RGBColor(*color)
    # bottom border
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'6')
    bot.set(qn('w:space'),'1'); bot.set(qn('w:color'),'%02X%02X%02X'%color)
    pb.append(bot); pPr.append(pb)
    return p

def h2(doc, text, color=(30,58,138)):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run(text); r.bold=True
    r.font.size=Pt(12); r.font.color.rgb=RGBColor(*color)
    return p

def step(doc, num, text, detail=''):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(f'  {num}.  ')
    r1.bold=True; r1.font.color.rgb=RGBColor(79,70,229); r1.font.size=Pt(12)
    r2 = p.add_run(text)
    r2.bold=True; r2.font.size=Pt(11)
    if detail:
        p2 = doc.add_paragraph(detail)
        p2.paragraph_format.left_indent = Inches(0.6)
        p2.paragraph_format.space_after = Pt(2)
        for r in p2.runs: r.font.size=Pt(10); r.font.color.rgb=RGBColor(71,85,105)
    return p

def info_box(doc, title, text, bg='EEF2FF', title_color=(79,70,229), icon='ℹ️'):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, bg)
    cell.width = Inches(5.5)
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_before = Pt(4)
    r = p1.add_run(f'{icon}  {title}')
    r.bold=True; r.font.size=Pt(11); r.font.color.rgb=RGBColor(*title_color)
    p2 = cell.add_paragraph(text)
    p2.paragraph_format.space_after = Pt(4)
    for r2 in p2.runs: r2.font.size=Pt(10)
    doc.add_paragraph()

def screen_box(doc, title, lines, color='4F46E5'):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'
    cell = tbl.cell(0,0)
    # Header
    ph = cell.paragraphs[0]
    ph.paragraph_format.space_before = Pt(2)
    rh = ph.add_run(f'  📱  {title}')
    rh.bold=True; rh.font.size=Pt(10); rh.font.color.rgb=RGBColor(255,255,255)
    set_cell_bg(cell, color)
    for line in lines:
        p = cell.add_paragraph(f'  {line}')
        p.paragraph_format.space_after = Pt(1)
        for r in p.runs: r.font.size=Pt(9); r.font.color.rgb=RGBColor(30,30,30)
    set_cell_bg(cell, 'F8FAFC')
    # Fix header bg
    doc.add_paragraph()

# ═══════════════════════════════
# COVER PAGE
# ═══════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
title_para(doc, '📋', 48, (79,70,229))
title_para(doc, 'BUKU PANDUAN PENGGUNAAN', 22, (79,70,229))
title_para(doc, 'Aplikasi Lost & Found', 28, (67,56,202))
title_para(doc, 'BXRink Ice Skating', 18, (100,116,139))
doc.add_paragraph()
title_para(doc, '─────────────────────────────', 14, (199,210,254))
doc.add_paragraph()
title_para(doc, 'Panduan Lengkap untuk Pengguna', 13, (107,114,128), bold=False)
title_para(doc, 'Mudah dipahami · Langkah demi langkah', 11, (148,163,184), bold=False)
doc.add_paragraph()
doc.add_paragraph()
title_para(doc, '🌐  stirring-yeot-a170188.netlify.app', 11, (99,102,241))
title_para(doc, '📞  +62 812-3456-7890 (WhatsApp)', 11, (22,163,74))
doc.add_page_break()

# ═══════════════════════════════
# DAFTAR ISI
# ═══════════════════════════════
title_para(doc, 'DAFTAR ISI', 16, (79,70,229), align=WD_ALIGN_PARAGRAPH.LEFT)
isi = [
    ('A', 'Mengenal Aplikasi', '3'),
    ('B', 'Cara Membuka Aplikasi', '3'),
    ('C', 'Untuk Pengunjung (Guest) — Lapor Barang Hilang', '4'),
    ('D', 'Untuk Pengunjung (Guest) — Cari Barang Ditemukan', '5'),
    ('E', 'Untuk Pengunjung (Guest) — Cara Klaim Barang', '5'),
    ('F', 'Untuk Rink Operator — Login', '6'),
    ('G', 'Untuk Rink Operator — Lapor Barang Ditemukan', '7'),
    ('H', 'Untuk Rink Operator — Proses Klaim Pengambilan', '8'),
    ('I', 'Untuk Rink Operator — Konfirmasi Serah Terima', '10'),
    ('J', 'Untuk Admin — Fitur Panel Admin', '11'),
    ('K', 'Informasi Penting', '13'),
]
tbl = doc.add_table(rows=len(isi), cols=3)
for i,(no,item,page) in enumerate(isi):
    row = tbl.rows[i]
    row.cells[0].text = no
    row.cells[1].text = item
    row.cells[2].text = page
    for j,cell in enumerate(row.cells):
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(11)
                if j==0: r.bold=True; r.font.color.rgb=RGBColor(79,70,229)
doc.add_page_break()

# ═══════════════════════════════
# A. MENGENAL APLIKASI
# ═══════════════════════════════
h1(doc, 'A', 'Mengenal Aplikasi Lost & Found BXRink')
p = doc.add_paragraph('Aplikasi Lost & Found adalah sistem digital untuk melaporkan, mencari, dan mengklaim barang yang hilang atau ditemukan di area BXRink Ice Skating.')
for r in p.runs: r.font.size=Pt(11)

doc.add_paragraph()
h2(doc, '🎯 Siapa yang bisa menggunakan aplikasi ini?')
tbl = doc.add_table(rows=4, cols=3)
tbl.style = 'Table Grid'
headers = ['Tipe Pengguna', 'Siapa', 'Yang Bisa Dilakukan']
for i, h in enumerate(headers):
    cell = tbl.rows[0].cells[i]
    cell.text = h
    set_cell_bg(cell, '4F46E5')
    for p in cell.paragraphs:
        for r in p.runs: r.bold=True; r.font.color.rgb=RGBColor(255,255,255); r.font.size=Pt(10)

rows_data = [
    ('👤 Guest', 'Pengunjung umum\n(tanpa login)', 'Lapor barang hilang\nCari barang ditemukan\nLihat panduan klaim'),
    ('🙋 Rink Operator', 'Staff BXRink\n(dengan login)', 'Semua fitur Guest +\nLapor barang ditemukan\nProses klaim & serah terima'),
    ('👑 Admin', 'Pengelola sistem\n(login superadmin)', 'Semua fitur Operator +\nPanel Admin, Export data\nPemusnahan & donasi'),
]
colors = ['F0F9FF','F5F3FF','EFF6FF']
for i,(a,b_,c) in enumerate(rows_data):
    row = tbl.rows[i+1]
    row.cells[0].text=a; row.cells[1].text=b_; row.cells[2].text=c
    for cell in row.cells:
        set_cell_bg(cell, colors[i])
        for p in cell.paragraphs:
            for r in p.runs: r.font.size=Pt(10)
doc.add_paragraph()
doc.add_page_break()

# ═══════════════════════════════
# B. CARA MEMBUKA APLIKASI
# ═══════════════════════════════
h1(doc, 'B', 'Cara Membuka Aplikasi')
step(doc, 1, 'Buka browser (Chrome, Firefox, atau Safari)')
step(doc, 2, 'Ketik alamat berikut di kolom URL:', 'stirring-yeot-a170188.netlify.app')
step(doc, 3, 'Tekan Enter — aplikasi langsung terbuka')
doc.add_paragraph()
info_box(doc, 'Tampilan Halaman Utama',
    'Saat pertama dibuka, Anda akan melihat halaman Barang Hilang dengan navbar di atas yang berisi menu:\n🚨 Barang Hilang  |  ✅ Barang Ditemukan  |  🤝 Pengambilan  |  👤 Masuk',
    'E0E7FF', (79,70,229), '🏠')

doc.add_paragraph()

# ═══════════════════════════════
# C. LAPOR BARANG HILANG (GUEST)
# ═══════════════════════════════
h1(doc, 'C', 'Lapor Barang Hilang (Pengunjung — Tanpa Login)')
p = doc.add_paragraph('Jika Anda kehilangan barang di area BXRink, ikuti langkah berikut:')
for r in p.runs: r.font.size=Pt(11)
doc.add_paragraph()

step(doc, 1, 'Klik menu 🚨 Barang Hilang di bagian atas halaman')
step(doc, 2, 'Klik tombol + Laporkan Hilang (warna merah, kanan atas)')
step(doc, 3, 'Isi formulir yang muncul:')

tbl = doc.add_table(rows=8, cols=2)
tbl.style = 'Table Grid'
fields = [
    ('Nama Barang *', 'Contoh: Sepatu Bauer, Tas Hitam, HP Samsung'),
    ('Kategori *', 'Pilih dari daftar: Blade Guard, Hockey Gear, Elektronik, dll'),
    ('Deskripsi *', 'Jelaskan ciri-ciri: warna, ukuran, merek, kondisi, isi'),
    ('Lokasi Kejadian *', 'Pilih dari daftar area BXRink'),
    ('Tanggal *', 'Tanggal barang hilang'),
    ('Nama Pelapor *', 'Nama lengkap Anda'),
    ('Kontak (WA) *', 'Nomor WhatsApp aktif untuk dihubungi'),
]
set_cell_bg(tbl.rows[0].cells[0], '4F46E5'); set_cell_bg(tbl.rows[0].cells[1], '4F46E5')
tbl.rows[0].cells[0].text = 'Field'; tbl.rows[0].cells[1].text = 'Keterangan'
for cell in tbl.rows[0].cells:
    for p in cell.paragraphs:
        for r in p.runs: r.bold=True; r.font.color.rgb=RGBColor(255,255,255); r.font.size=Pt(10)
for i,(f,k) in enumerate(fields):
    tbl.rows[i+1].cells[0].text=f; tbl.rows[i+1].cells[1].text=k
    for cell in tbl.rows[i+1].cells:
        set_cell_bg(cell, 'FEF2F2' if i%2==0 else 'FFFFFF')
        for p in cell.paragraphs:
            for r in p.runs: r.font.size=Pt(10)
doc.add_paragraph()
step(doc, 4, 'Upload foto barang (opsional tapi sangat disarankan)', 'Klik area foto atau seret foto ke kotak yang tersedia · Format: JPG/PNG · Maks 5 MB')
step(doc, 5, 'Centang reCAPTCHA "I am not a robot"')
step(doc, 6, 'Klik Kirim Laporan')
info_box(doc, 'Berhasil!', 'Laporan Anda tersimpan. Sistem akan otomatis mencocokkan dengan barang ditemukan yang ada.', 'F0FDF4', (22,163,74), '✅')
doc.add_page_break()

# ═══════════════════════════════
# D. CARI BARANG (GUEST)
# ═══════════════════════════════
h1(doc, 'D', 'Mencari Barang yang Ditemukan (Pengunjung)')
p = doc.add_paragraph('Cek apakah barang Anda sudah ditemukan dan dilaporkan di aplikasi:')
for r in p.runs: r.font.size=Pt(11)
doc.add_paragraph()
step(doc, 1, 'Klik menu ✅ Barang Ditemukan di navbar')
step(doc, 2, 'Ketik nama atau deskripsi barang di kolom pencarian', 'Contoh: "sepatu bauer" atau "tas hitam hockey"')
step(doc, 3, 'Tekan Enter atau tunggu hasil muncul')
step(doc, 4, 'Scroll dan cari barang yang sesuai')
step(doc, 5, 'Klik kartu barang untuk melihat detail')
doc.add_paragraph()
info_box(doc, 'Tips Pencarian',
    '• Gunakan kata kunci singkat: "sepatu", "tas hitam", "kunci"\n• Filter berdasarkan kategori menggunakan dropdown\n• Jika barang ada, klik kartu untuk melihat detail lengkap\n• Foto barang hanya terlihat setelah login',
    'FFF7ED', (146,64,14), '💡')

doc.add_paragraph()
# ═══════════════════════════════
# E. CARA KLAIM (GUEST)
# ═══════════════════════════════
h1(doc, 'E', 'Cara Klaim Barang — untuk Pengunjung')
p = doc.add_paragraph('Karena pengunjung tidak memiliki akun, proses klaim dilakukan melalui petugas BXRink:')
for r in p.runs: r.font.size=Pt(11)
doc.add_paragraph()
step(doc, 1, 'Cari barang Anda di menu ✅ Barang Ditemukan')
step(doc, 2, 'Hubungi petugas BXRink', '📱 WhatsApp: +62 812-3456-7890\n📞 Telepon: +62 812-3456-7890\n🕙 Jam: Senin–Minggu, 10.00–17.00 WIB')
step(doc, 3, 'Jelaskan ciri-ciri barang Anda secara detail ke petugas', 'Warna, merek, isi, ciri khas yang hanya pemilik tahu')
step(doc, 4, 'Petugas akan jadwalkan waktu pengambilan barang di BXRink')
step(doc, 5, 'Datang ke BXRink sesuai jadwal untuk mengambil barang')
doc.add_page_break()

# ═══════════════════════════════
# F. LOGIN RINK OPERATOR
# ═══════════════════════════════
h1(doc, 'F', 'Login sebagai Rink Operator', (37,99,235))
p = doc.add_paragraph('Staff BXRink (Rink Operator) perlu login untuk mengakses fitur tambahan:')
for r in p.runs: r.font.size=Pt(11)
doc.add_paragraph()
step(doc, 1, 'Klik tombol 👤 Masuk di pojok kanan atas navbar')
step(doc, 2, 'Muncul form login — isi Username dan Password:')
tbl = doc.add_table(rows=3, cols=2); tbl.style='Table Grid'
set_cell_bg(tbl.rows[0].cells[0],'2563EB'); set_cell_bg(tbl.rows[0].cells[1],'2563EB')
tbl.rows[0].cells[0].text='Field'; tbl.rows[0].cells[1].text='Yang Harus Diisi'
for cell in tbl.rows[0].cells:
    for p in cell.paragraphs:
        for r in p.runs: r.bold=True; r.font.color.rgb=RGBColor(255,255,255); r.font.size=Pt(10)
data=[('Username','rinkoperator'),('Password','(password diberikan pengelola)')]
for i,(f,v) in enumerate(data):
    tbl.rows[i+1].cells[0].text=f; tbl.rows[i+1].cells[1].text=v
    set_cell_bg(tbl.rows[i+1].cells[0],'EFF6FF'); set_cell_bg(tbl.rows[i+1].cells[1],'FFFFFF')
    for cell in tbl.rows[i+1].cells:
        for p in cell.paragraphs:
            for r in p.runs: r.font.size=Pt(11)
doc.add_paragraph()
step(doc, 3, 'Klik Masuk')
step(doc, 4, 'Setelah berhasil, nama Anda muncul di navbar: Rink Operator 🙋')
info_box(doc,'Setelah Login','Menu tambahan akan muncul:\n• + Laporkan Temuan (di halaman Barang Ditemukan)\n• Tab Ajukan Klaim & Cek Status (di menu Pengambilan)\n• Lihat foto dan kontak pelapor di detail barang','EFF6FF',(37,99,235),'✅')
doc.add_paragraph()

# ═══════════════════════════════
# G. LAPOR BARANG DITEMUKAN
# ═══════════════════════════════
h1(doc, 'G', 'Lapor Barang Ditemukan (Rink Operator)', (37,99,235))
p = doc.add_paragraph('Jika menemukan barang tidak bertuan di area BXRink:')
for r in p.runs: r.font.size=Pt(11)
doc.add_paragraph()
step(doc, 1, 'Pastikan sudah login sebagai Rink Operator')
step(doc, 2, 'Klik menu ✅ Barang Ditemukan di navbar')
step(doc, 3, 'Klik tombol + Laporkan Temuan (warna hijau, kanan atas)')
step(doc, 4, 'Isi formulir:')
tbl2 = doc.add_table(rows=9, cols=2); tbl2.style='Table Grid'
set_cell_bg(tbl2.rows[0].cells[0],'16A34A'); set_cell_bg(tbl2.rows[0].cells[1],'16A34A')
tbl2.rows[0].cells[0].text='Field'; tbl2.rows[0].cells[1].text='Keterangan'
for cell in tbl2.rows[0].cells:
    for p in cell.paragraphs:
        for r in p.runs: r.bold=True; r.font.color.rgb=RGBColor(255,255,255); r.font.size=Pt(10)
f2=[('Nama Barang *','Nama/jenis barang yang ditemukan'),('Kategori *','Blade Guard, Hockey Gear, Elektronik, dll'),('Deskripsi *','Ciri-ciri: warna, kondisi, isi, merek'),('Lokasi *','Area BXRink tempat barang ditemukan'),('Tanggal *','Tanggal barang ditemukan'),('Nama Penemu *','Nama staff yang menemukan'),('Kontak *','Nomor WA untuk dihubungi'),('Foto','Upload foto barang (sangat disarankan)')]
for i,(f,k) in enumerate(f2):
    tbl2.rows[i+1].cells[0].text=f; tbl2.rows[i+1].cells[1].text=k
    set_cell_bg(tbl2.rows[i+1].cells[0],'F0FDF4' if i%2==0 else 'FFFFFF')
    for cell in tbl2.rows[i+1].cells:
        for p in cell.paragraphs:
            for r in p.runs: r.font.size=Pt(10)
doc.add_paragraph()
step(doc, 5, 'Klik Kirim Laporan')
info_box(doc,'Sistem Auto-Match','Setelah laporan dikirim, sistem otomatis mencari barang hilang yang mungkin cocok berdasarkan deskripsi dan foto. Jika ditemukan kandidat cocok, akan muncul panel perbandingan.','F0FDF4',(22,163,74),'🔍')
doc.add_page_break()

# ═══════════════════════════════
# H. PROSES KLAIM
# ═══════════════════════════════
h1(doc, 'H', 'Proses Klaim Pengambilan (Rink Operator)', (37,99,235))
p = doc.add_paragraph('Ketika pemilik barang datang untuk mengambil barang, Rink Operator memproses klaim:')
for r in p.runs: r.font.size=Pt(11)
doc.add_paragraph()
info_box(doc,'Penting!','Klaim langsung selesai setelah disubmit — TIDAK perlu menunggu verifikasi admin. Barang otomatis hilang dari daftar.','FEF9C3',(146,64,14),'⚡')
doc.add_paragraph()

h2(doc,'📝 Step 1: Pilih Barang',(37,99,235))
step(doc,1,'Buka menu 🤝 Pengambilan di navbar')
step(doc,2,'Klik tab Ajukan Klaim')
step(doc,3,'Cari barang yang ingin diklaim menggunakan kolom pencarian')
step(doc,4,'Klik pada kartu barang untuk memilihnya','Tanda barang dipilih: border berubah warna')
step(doc,5,'Klik tombol Lanjut →')

doc.add_paragraph()
h2(doc,'📋 Step 2: Bukti Kepemilikan',(37,99,235))
step(doc,1,'Isi Nama Pengklaim','Nama lengkap pemilik barang')
step(doc,2,'Isi Nomor Kontak (WA)','Nomor WhatsApp pemilik barang')
step(doc,3,'Isi Nama Petugas','Nama staff BXRink yang bertugas')
step(doc,4,'Isi Bukti Kepemilikan','Minta pemilik mendeskripsikan detail barang yang hanya pemilik asli yang tahu\nContoh: "Di dalam tas ada charger laptop merek Dell dan buku catatan berwarna merah"')
step(doc,5,'Upload Foto Pengambil','Klik area foto → pilih foto pemilik barang → foto terupload sebagai bukti')
step(doc,6,'Klik Lanjut →')

doc.add_paragraph()
h2(doc,'📅 Step 3: Jadwal Pengambilan',(37,99,235))
step(doc,1,'Isi Tanggal Pengambilan Barang','Tanggal hari ini atau sesuai kesepakatan')
step(doc,2,'Pilih Waktu','Pilih rentang waktu pengambilan (08:00 – 22:00)')
step(doc,3,'Isi Catatan Tambahan (opsional)','Keterangan tambahan jika diperlukan')
step(doc,4,'Klik Kirim Klaim ✅')
doc.add_paragraph()
info_box(doc,'Klaim Berhasil!','✅ Status klaim: Selesai\n✅ Barang otomatis hilang dari daftar\n✅ Data tersimpan di riwayat pengambilan\n✅ Data tersync ke Google Sheets otomatis','F0FDF4',(22,163,74),'🎉')
doc.add_page_break()

# ═══════════════════════════════
# I. KONFIRMASI SERAH TERIMA
# ═══════════════════════════════
h1(doc,'I','Konfirmasi Serah Terima (Rink Operator)',(37,99,235))
p=doc.add_paragraph('Setelah barang diserahkan ke pemilik, upload foto bukti serah terima:')
for r in p.runs: r.font.size=Pt(11)
doc.add_paragraph()
step(doc,1,'Buka menu 🤝 Pengambilan → tab Cek Status')
step(doc,2,'Cari klaim yang sudah selesai (masukkan nama atau nomor kontak pemilik)')
step(doc,3,'Akan muncul panel biru "📋 Konfirmasi Serah Terima"')
step(doc,4,'Klik tombol 📸 Upload Bukti Serah Terima')
step(doc,5,'Upload foto','Foto pemilik menerima barang atau foto barang saat diserahkan')
step(doc,6,'Isi catatan serah terima (opsional)')
step(doc,7,'Klik Konfirmasi Selesai')
doc.add_paragraph()
info_box(doc,'Selesai!','Bukti foto tersimpan permanen di sistem dan dapat dilihat oleh Admin di Riwayat Pengambilan.','F0FDF4',(22,163,74),'✅')
doc.add_page_break()

# ═══════════════════════════════
# J. PANEL ADMIN
# ═══════════════════════════════
h1(doc,'J','Fitur Panel Admin (Superadmin)',(79,70,229))
p=doc.add_paragraph('Login sebagai superadmin untuk mengakses fitur pengelolaan lengkap:')
for r in p.runs: r.font.size=Pt(11)
doc.add_paragraph()

h2(doc,'🔐 Cara Login Admin',(79,70,229))
step(doc,1,'Tekan Ctrl+Shift+A di keyboard, ATAU klik logo "🔍 Lost & Found" sebanyak 5 kali')
step(doc,2,'Halaman Admin terbuka → masukkan Username: superadmin dan Password')
step(doc,3,'Klik Masuk → Menu admin muncul di navbar')
doc.add_paragraph()

h2(doc,'📊 Dashboard',(79,70,229))
p=doc.add_paragraph('Buka tab 📊 Dashboard untuk melihat statistik lengkap:')
for r in p.runs: r.font.size=Pt(11)
step(doc,'•','Grafik laporan per bulan (barang hilang & ditemukan)')
step(doc,'•','Statistik status laporan (aktif, selesai, dimusnahkan)')
step(doc,'•','Breakdown kategori barang')
step(doc,'•','Lokasi dengan kejadian terbanyak')
doc.add_paragraph()

h2(doc,'📋 Semua Laporan',(79,70,229))
step(doc,'•','Lihat timeline semua laporan dari seluruh pengguna')
step(doc,'•','Filter berdasarkan status, kategori, atau tanggal')
step(doc,'•','Konfirmasi kecocokan barang hilang & ditemukan')
doc.add_paragraph()

h2(doc,'⬇️ Export Data',(79,70,229))
step(doc,1,'Di tab Barang Hilang, klik tombol ⬇️ Export CSV')
step(doc,2,'File Excel otomatis terdownload ke komputer')
step(doc,3,'Tersedia untuk: Barang Hilang, Barang Ditemukan, Data Klaim')
info_box(doc,'Google Sheets Sync','Data otomatis tersync ke Google Sheets setiap ada perubahan. Tidak perlu export manual untuk melihat data terbaru di Sheets.','EEF2FF',(79,70,229),'📊')
doc.add_paragraph()

h2(doc,'🗑️ Pemusnahan / Donasi Barang',(220,38,38))
p=doc.add_paragraph('Untuk barang yang tidak diklaim lebih dari 30 hari:')
for r in p.runs: r.font.size=Pt(11)
step(doc,1,'Buka menu 🗑️ Pemusnahan di navbar')
step(doc,2,'Daftar barang eligible akan muncul (belum ada klaim, lebih dari 30 hari)')
step(doc,3,'Pilih barang → klik Proses Pemusnahan / Donasi')
step(doc,4,'Pilih tindakan: 🔥 Dimusnahkan atau 💙 Didonasikan')
step(doc,5,'Isi formulir lengkap dan upload foto bukti')
step(doc,6,'Centang pernyataan → klik Proses Tindakan')
doc.add_page_break()

# ═══════════════════════════════
# K. INFORMASI PENTING
# ═══════════════════════════════
h1(doc,'K','Informasi Penting')

h2(doc,'📍 Area Lokasi BXRink',(8,145,178))
locs=['Ice Rink Arena','Member Room Figure','Member Room Hockey','Multifunction Room / Off Ice Room','Waiting Area','Prayer Room','Toilet','Other']
for l in locs: step(doc,'•',l)
doc.add_paragraph()

h2(doc,'📁 Kategori Barang',(8,145,178))
cats=['Blade Guard','Botol Minum','Dompet / Tas','Elektronik','Hockey Gear','Kunci','Pakaian','Lainnya']
p=doc.add_paragraph(' · '.join(cats))
for r in p.runs: r.font.size=Pt(11)
doc.add_paragraph()

h2(doc,'📞 Hubungi Kami',(22,163,74))
step(doc,'📱','WhatsApp: +62 812-3456-7890')
step(doc,'🕙','Jam Operasional: Senin – Minggu, 10.00 – 17.00 WIB')
step(doc,'🌐','Aplikasi: stirring-yeot-a170188.netlify.app')
doc.add_paragraph()

info_box(doc,'Pertanyaan Umum',
    'T: Barang sudah saya laporkan tapi belum ditemukan, apa yang harus dilakukan?\nJ: Tunggu beberapa waktu dan cek berkala di aplikasi. Jika ada barang ditemukan yang cocok, sistem akan mencocokkan otomatis.\n\nT: Berapa lama barang disimpan?\nJ: Barang disimpan hingga diklaim. Setelah 30 hari tidak ada klaim, barang dapat dimusnahkan atau didonasikan sesuai prosedur.\n\nT: Apakah data saya aman?\nJ: Ya, data tersimpan di Firebase Google yang terenkripsi dan aman.',
    'F0F9FF',(8,145,178),'❓')

# FOOTER
doc.add_paragraph()
p=doc.add_paragraph('Aplikasi Lost & Found BXRink — Dibuat untuk kemudahan pengelolaan barang hilang & ditemukan')
p.alignment=WD_ALIGN_PARAGRAPH.CENTER
for r in p.runs: r.font.size=Pt(9); r.font.color.rgb=RGBColor(148,163,184)

out=r'C:\Users\Lenovo\OneDrive\Documents\lostandfound\Manual_Book_LostFound.docx'
doc.save(out)
print('SUCCESS:', out)
