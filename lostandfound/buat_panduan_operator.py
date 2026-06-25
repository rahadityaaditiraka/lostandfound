from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
s = doc.sections[0]
s.page_width  = Inches(8.27)
s.page_height = Inches(11.69)
s.top_margin = Inches(0.9)
s.bottom_margin = Inches(0.9)
s.left_margin = Inches(1.1)
s.right_margin = Inches(1.1)

# ── Helpers ────────────────────────────────────
def cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def para(doc, text='', size=11, bold=False, color=(30,30,30), align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6):
    p = doc.add_paragraph(); p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        r = p.add_run(text); r.bold=bold
        r.font.size=Pt(size); r.font.color.rgb=RGBColor(*color)
    return p

def section_header(doc, number, title, color='2563EB'):
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0,0)
    cell_bg(cell, color)
    cell.width = Inches(6.07)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    r1 = p.add_run(f'  {number}  '); r1.bold=True; r1.font.size=Pt(14); r1.font.color.rgb=RGBColor(255,255,255)
    r2 = p.add_run(title); r2.bold=True; r2.font.size=Pt(13); r2.font.color.rgb=RGBColor(255,255,255)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def step_box(doc, number, title, details=None, color='1E3A8A'):
    tbl = doc.add_table(rows=1, cols=2)
    tbl.columns[0].width = Inches(0.45)
    tbl.columns[1].width = Inches(5.62)
    # Number circle
    c0 = tbl.cell(0,0); cell_bg(c0, color)
    p0 = c0.paragraphs[0]; p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0.paragraph_format.space_before = Pt(6)
    p0.paragraph_format.space_after  = Pt(6)
    r0 = p0.add_run(str(number)); r0.bold=True; r0.font.size=Pt(14); r0.font.color.rgb=RGBColor(255,255,255)
    # Content
    c1 = tbl.cell(0,1); cell_bg(c1, 'F0F9FF')
    p1 = c1.paragraphs[0]
    p1.paragraph_format.space_before = Pt(5)
    r1 = p1.add_run(title); r1.bold=True; r1.font.size=Pt(11); r1.font.color.rgb=RGBColor(30,64,175)
    if details:
        for d in details:
            p2 = c1.add_paragraph(f'  → {d}')
            p2.paragraph_format.space_after = Pt(1)
            for r in p2.runs: r.font.size=Pt(10); r.font.color.rgb=RGBColor(71,85,105)
    c1.paragraphs[-1].paragraph_format.space_after = Pt(5)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)

def note_box(doc, emoji, title, body, bg='FEF9C3', title_rgb=(146,64,14)):
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0,0); cell_bg(cell, bg)
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_before = Pt(6)
    r1 = p1.add_run(f'  {emoji}  {title}')
    r1.bold=True; r1.font.size=Pt(11); r1.font.color.rgb=RGBColor(*title_rgb)
    for line in body.split('\n'):
        p2 = cell.add_paragraph(f'  {line}')
        for r in p2.runs: r.font.size=Pt(10)
    cell.paragraphs[-1].paragraph_format.space_after = Pt(6)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def field_row(tbl, row_idx, label, value, bg_label='DBEAFE', bg_val='EFF6FF'):
    row = tbl.rows[row_idx]
    row.cells[0].text = label; row.cells[1].text = value
    cell_bg(row.cells[0], bg_label); cell_bg(row.cells[1], bg_val)
    for cell in row.cells:
        for p in cell.paragraphs:
            for r in p.runs: r.font.size=Pt(10)
    row.cells[0].paragraphs[0].runs[0].bold = True

# ═══════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════
para(doc, space_before=20)
para(doc, space_before=20)

tbl_cover = doc.add_table(rows=1, cols=1)
cell_cover = tbl_cover.cell(0,0)
cell_bg(cell_cover, '2563EB')

pc1 = cell_cover.paragraphs[0]; pc1.alignment = WD_ALIGN_PARAGRAPH.CENTER
pc1.paragraph_format.space_before = Pt(24)
r = pc1.add_run('🙋'); r.font.size = Pt(48)

pc2 = cell_cover.add_paragraph(); pc2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = pc2.add_run('PANDUAN PENGGUNAAN'); r.bold=True; r.font.size=Pt(22); r.font.color.rgb=RGBColor(255,255,255)

pc3 = cell_cover.add_paragraph(); pc3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = pc3.add_run('Aplikasi Lost & Found BXRink'); r.bold=True; r.font.size=Pt(16); r.font.color.rgb=RGBColor(191,219,254)

pc4 = cell_cover.add_paragraph(); pc4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = pc4.add_run('Khusus Rink Operator'); r.bold=True; r.font.size=Pt(13); r.font.color.rgb=RGBColor(147,197,253)

pc5 = cell_cover.add_paragraph(); pc5.alignment = WD_ALIGN_PARAGRAPH.CENTER
pc5.paragraph_format.space_before = Pt(16)
r = pc5.add_run('━' * 30); r.font.color.rgb=RGBColor(96,165,250); r.font.size=Pt(10)

pc6 = cell_cover.add_paragraph(); pc6.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = pc6.add_run('BXRink Ice Skating'); r.font.size=Pt(12); r.font.color.rgb=RGBColor(191,219,254)

pc7 = cell_cover.add_paragraph(); pc7.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = pc7.add_run('stirring-yeot-a170188.netlify.app'); r.font.size=Pt(11); r.font.color.rgb=RGBColor(147,197,253)
pc7.paragraph_format.space_after = Pt(24)

doc.add_paragraph()
# Credential box
tbl_cred = doc.add_table(rows=3, cols=2); tbl_cred.style='Table Grid'
headers = ['Informasi Login', '']
cell_bg(tbl_cred.rows[0].cells[0], '1E40AF'); cell_bg(tbl_cred.rows[0].cells[1], '1E40AF')
tbl_cred.rows[0].cells[0].merge(tbl_cred.rows[0].cells[1])
tbl_cred.rows[0].cells[0].text = '  🔐  Informasi Login Rink Operator'
for r in tbl_cred.rows[0].cells[0].paragraphs[0].runs:
    r.bold=True; r.font.color.rgb=RGBColor(255,255,255); r.font.size=Pt(11)

data = [('  Username','  rinkoperator'), ('  Password','  (diberikan oleh pengelola)')]
for i,(a,b) in enumerate(data):
    tbl_cred.rows[i+1].cells[0].text=a; tbl_cred.rows[i+1].cells[1].text=b
    cell_bg(tbl_cred.rows[i+1].cells[0],'DBEAFE'); cell_bg(tbl_cred.rows[i+1].cells[1],'EFF6FF')
    for cell in tbl_cred.rows[i+1].cells:
        for p in cell.paragraphs:
            for r in p.runs: r.font.size=Pt(11); r.bold=(tbl_cred.rows[i+1].cells[0]==cell)

doc.add_page_break()

# ═══════════════════════════════════════════════
# DAFTAR ISI
# ═══════════════════════════════════════════════
para(doc,'DAFTAR ISI',14,True,(37,99,235),WD_ALIGN_PARAGRAPH.LEFT,0,8)
isi = [
    ('1','Cara Login ke Aplikasi','3'),
    ('2','Cara Lapor Barang Ditemukan','4'),
    ('3','Cara Melihat Detail Barang','5'),
    ('4','Proses Klaim Pengambilan Barang','6'),
    ('  Step 1','— Pilih Barang','7'),
    ('  Step 2','— Bukti Kepemilikan','7'),
    ('  Step 3','— Jadwal Pengambilan','8'),
    ('5','Konfirmasi Serah Terima Barang','9'),
    ('6','Cara Cek Status Klaim','10'),
    ('7','Cara Logout','10'),
]
tbl_isi = doc.add_table(rows=len(isi), cols=3)
for i,(no,item,pg) in enumerate(isi):
    tbl_isi.rows[i].cells[0].text=no
    tbl_isi.rows[i].cells[1].text=item
    tbl_isi.rows[i].cells[2].text=pg
    for j,cell in enumerate(tbl_isi.rows[i].cells):
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size=Pt(11)
                if j==0 and not no.startswith(' '): r.bold=True; r.font.color.rgb=RGBColor(37,99,235)
doc.add_page_break()

# ═══════════════════════════════════════════════
# 1. CARA LOGIN
# ═══════════════════════════════════════════════
section_header(doc,'1','Cara Login ke Aplikasi')

para(doc,'Sebelum menggunakan fitur operator, Anda harus login terlebih dahulu.',11,space_after=8)

step_box(doc, 1, 'Buka aplikasi di browser',
    ['Ketik di address bar: stirring-yeot-a170188.netlify.app', 'Tekan Enter'])
step_box(doc, 2, 'Klik tombol 👤 Masuk',
    ['Tombol ada di pojok kanan atas layar (navbar)'])
step_box(doc, 3, 'Isi Username dan Password',
    ['Username: rinkoperator', 'Password: (sesuai yang diberikan pengelola)'])
step_box(doc, 4, 'Centang reCAPTCHA',
    ['Klik kotak "I am not a robot"', 'Ikuti instruksi jika diminta memilih gambar'])
step_box(doc, 5, 'Klik tombol Masuk',
    ['Jika berhasil, nama "Rink Operator 🙋" akan muncul di kanan atas'])

note_box(doc,'✅','Login Berhasil!',
    'Setelah login berhasil, Anda akan melihat tambahan fitur:\n'
    '• Tombol + Laporkan Temuan di halaman Barang Ditemukan\n'
    '• Tab "Ajukan Klaim" dan "Cek Status" di menu Pengambilan\n'
    '• Foto & kontak pelapor bisa dilihat di detail barang',
    'F0FDF4',(22,163,74))

note_box(doc,'⚠️','Jika Login Gagal',
    'Periksa kembali Username dan Password — perhatikan huruf kapital\n'
    'Pastikan reCAPTCHA sudah dicentang\n'
    'Hubungi Admin jika masalah berlanjut: +62 812-3456-7890',
    'FEF2F2',(220,38,38))

doc.add_page_break()

# ═══════════════════════════════════════════════
# 2. LAPOR BARANG DITEMUKAN
# ═══════════════════════════════════════════════
section_header(doc,'2','Cara Lapor Barang Ditemukan')

para(doc,'Jika Anda menemukan barang tidak bertuan di area BXRink, segera laporkan ke aplikasi:',11,space_after=8)

step_box(doc, 1, 'Klik menu ✅ Barang Ditemukan di navbar')
step_box(doc, 2, 'Klik tombol + Laporkan Temuan (warna hijau, kanan atas)')
step_box(doc, 3, 'Isi formulir laporan dengan lengkap:')

tbl_form = doc.add_table(rows=9, cols=2); tbl_form.style='Table Grid'
cell_bg(tbl_form.rows[0].cells[0],'1E40AF'); cell_bg(tbl_form.rows[0].cells[1],'1E40AF')
tbl_form.rows[0].cells[0].text='  Field'; tbl_form.rows[0].cells[1].text='  Cara Mengisi'
for cell in tbl_form.rows[0].cells:
    for p in cell.paragraphs:
        for r in p.runs: r.bold=True; r.font.color.rgb=RGBColor(255,255,255); r.font.size=Pt(10)

fields = [
    ('Nama Barang ★',    'Ketik nama barang  Contoh: Sepatu Bauer, Tas Hitam, Helm'),
    ('Kategori ★',       'Pilih dari daftar  Blade Guard, Hockey Gear, Elektronik, dll'),
    ('Deskripsi ★',      'Jelaskan detail: warna, merek, kondisi, isi atau ciri khas'),
    ('Lokasi ★',         'Pilih area tempat barang ditemukan dari daftar dropdown'),
    ('Tanggal ★',        'Tanggal barang ditemukan (isi tanggal hari ini jika baru)'),
    ('Nama Penemu ★',    'Nama staff / Anda yang menemukan barang'),
    ('Kontak (WA) ★',    'Nomor WhatsApp yang bisa dihubungi (hanya angka)'),
    ('Foto Barang',      'SANGAT DISARANKAN — klik area foto dan pilih foto barang'),
]
for i,(f,k) in enumerate(fields):
    bg0 = 'DBEAFE' if i%2==0 else 'EFF6FF'
    bg1 = 'F0F9FF' if i%2==0 else 'FFFFFF'
    tbl_form.rows[i+1].cells[0].text=f'  {f}'; tbl_form.rows[i+1].cells[1].text=f'  {k}'
    cell_bg(tbl_form.rows[i+1].cells[0],bg0); cell_bg(tbl_form.rows[i+1].cells[1],bg1)
    for cell in tbl_form.rows[i+1].cells:
        for p in cell.paragraphs:
            for r in p.runs: r.font.size=Pt(10)
    tbl_form.rows[i+1].cells[0].paragraphs[0].runs[0].bold=True

doc.add_paragraph()
step_box(doc, 4, 'Klik Kirim Laporan')

note_box(doc,'💡','Tips Pengisian Deskripsi',
    'Semakin detail deskripsi, semakin mudah pemilik membuktikan kepemilikannya.\n'
    'Contoh deskripsi bagus: "Tas ransel warna hitam merek Bauer. Berisi sepatu ice skating ukuran 40,\n'
    'sarung tangan biru, dan botol minum merah."',
    'EFF6FF',(37,99,235))

doc.add_page_break()

# ═══════════════════════════════════════════════
# 3. MELIHAT DETAIL BARANG
# ═══════════════════════════════════════════════
section_header(doc,'3','Cara Melihat Detail Barang')

para(doc,'Sebagai Rink Operator, Anda bisa melihat informasi lengkap termasuk foto dan kontak pelapor:',11,space_after=8)

step_box(doc, 1, 'Buka halaman Barang Hilang atau Barang Ditemukan')
step_box(doc, 2, 'Gunakan kolom pencarian untuk mencari barang tertentu',
    ['Ketik nama atau deskripsi barang', 'Contoh: "sepatu", "tas hitam", "blade guard"'])
step_box(doc, 3, 'Klik kartu barang untuk membuka detail')
step_box(doc, 4, 'Lihat informasi lengkap:',
    ['Foto barang (klik "🖼️ Lihat Barang" untuk tampilan penuh)',
     'Kategori, deskripsi, lokasi, tanggal',
     'Nama pelapor dan nomor kontak WhatsApp'])
step_box(doc, 5, 'Untuk menghubungi pelapor, klik nomor kontak yang muncul berwarna biru')

note_box(doc,'ℹ️','Informasi Penting',
    'Foto barang dan kontak pelapor HANYA terlihat setelah login sebagai Rink Operator atau Admin.\n'
    'Pengunjung biasa tidak bisa melihat informasi tersebut.',
    'EEF2FF',(79,70,229))

doc.add_page_break()

# ═══════════════════════════════════════════════
# 4. PROSES KLAIM
# ═══════════════════════════════════════════════
section_header(doc,'4','Proses Klaim Pengambilan Barang')

para(doc,'Ketika pemilik barang datang untuk mengambil barangnya, ikuti proses klaim 3 langkah ini:',11,space_after=4)

note_box(doc,'⚡','Perlu Diketahui',
    'Klaim LANGSUNG SELESAI setelah dikirim — tidak perlu menunggu konfirmasi admin.\n'
    'Barang otomatis hilang dari daftar setelah klaim berhasil.',
    'FEF9C3',(146,64,14))

# STEP 1
tbl_s1 = doc.add_table(rows=1,cols=1)
cell_s1 = tbl_s1.cell(0,0); cell_bg(cell_s1,'1E3A8A')
p = cell_s1.paragraphs[0]; p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(4)
r = p.add_run('  STEP 1 — Pilih Barang  '); r.bold=True; r.font.size=Pt(12); r.font.color.rgb=RGBColor(255,255,255)
doc.add_paragraph().paragraph_format.space_after=Pt(2)

step_box(doc,'a','Klik menu 🤝 Pengambilan di navbar atas')
step_box(doc,'b','Klik tab Ajukan Klaim')
step_box(doc,'c','Cari barang dengan mengetik nama di kolom pencarian')
step_box(doc,'d','Klik pada kartu barang yang ingin diklaim',
    ['Border kartu akan berubah warna menandakan barang terpilih',
     'Di bawah muncul teks "Dipilih: [nama barang]"'])
step_box(doc,'e','Klik tombol Lanjut →')

doc.add_paragraph()
# STEP 2
tbl_s2 = doc.add_table(rows=1,cols=1)
cell_s2 = tbl_s2.cell(0,0); cell_bg(cell_s2,'1E3A8A')
p2 = cell_s2.paragraphs[0]; p2.paragraph_format.space_before=Pt(4); p2.paragraph_format.space_after=Pt(4)
r2 = p2.add_run('  STEP 2 — Bukti Kepemilikan  '); r2.bold=True; r2.font.size=Pt(12); r2.font.color.rgb=RGBColor(255,255,255)
doc.add_paragraph().paragraph_format.space_after=Pt(2)

step_box(doc,'a','Isi Nama Pengklaim',['Nama lengkap pemilik barang yang datang mengambil'])
step_box(doc,'b','Isi Kontak (WA)','')
step_box(doc,'c','Isi Nama Petugas',['Nama staff BXRink yang bertugas menerima klaim'])
step_box(doc,'d','Isi Bukti Kepemilikan',
    ['Minta pemilik mendeskripsikan isi/ciri khas barang secara detail',
     'Contoh: "Di dalam tas ada charger laptop, buku catatan merah, dan kartu nama"',
     'Semakin detail semakin mudah diverifikasi'])
step_box(doc,'e','Upload Foto Pengambil',
    ['Klik area foto atau klik kotak "Klik atau seret foto ke sini"',
     'Pilih foto pemilik barang dari kamera / galeri',
     'Format: JPG atau PNG, maksimal 5 MB'])
step_box(doc,'f','Klik tombol Lanjut →')

doc.add_page_break()

# STEP 3
tbl_s3 = doc.add_table(rows=1,cols=1)
cell_s3 = tbl_s3.cell(0,0); cell_bg(cell_s3,'1E3A8A')
p3 = cell_s3.paragraphs[0]; p3.paragraph_format.space_before=Pt(4); p3.paragraph_format.space_after=Pt(4)
r3 = p3.add_run('  STEP 3 — Jadwal Pengambilan  '); r3.bold=True; r3.font.size=Pt(12); r3.font.color.rgb=RGBColor(255,255,255)
doc.add_paragraph().paragraph_format.space_after=Pt(2)

step_box(doc,'a','Isi Tanggal Pengambilan Barang',
    ['Pilih tanggal hari ini jika barang langsung diambil sekarang',
     'Atau pilih tanggal yang disepakati dengan pemilik'])
step_box(doc,'b','Pilih Waktu',
    ['Pilih rentang waktu dari dropdown (08:00 – 22:00)'])
step_box(doc,'c','Isi Catatan Tambahan (opsional)',
    ['Contoh: "Barang diserahkan langsung di Member Room Hockey"'])
step_box(doc,'d','Klik tombol Kirim Klaim ✅')

doc.add_paragraph()
note_box(doc,'🎉','Klaim Berhasil Dikirim!',
    'Setelah klik Kirim Klaim:\n'
    '✅ Status klaim berubah menjadi Selesai\n'
    '✅ Barang otomatis hilang dari daftar Barang Ditemukan\n'
    '✅ Data tersimpan di sistem dan Google Sheets\n'
    '✅ Anda akan diarahkan ke halaman Cek Status',
    'F0FDF4',(22,163,74))

doc.add_page_break()

# ═══════════════════════════════════════════════
# 5. KONFIRMASI SERAH TERIMA
# ═══════════════════════════════════════════════
section_header(doc,'5','Konfirmasi Serah Terima Barang')

para(doc,'Setelah barang diserahkan ke pemilik, upload foto bukti sebagai dokumentasi resmi:',11,space_after=8)

step_box(doc, 1, 'Klik menu 🤝 Pengambilan di navbar')
step_box(doc, 2, 'Klik tab Cek Status')
step_box(doc, 3, 'Masukkan nama atau nomor kontak pengklaim',['Klik tombol Cek'])
step_box(doc, 4, 'Temukan kartu klaim yang muncul')
step_box(doc, 5, 'Klik tombol 📸 Upload Bukti Serah Terima (panel biru di bawah kartu)')
step_box(doc, 6, 'Upload foto serah terima',
    ['Foto pemilik menerima barang, atau foto barang saat diserahkan',
     'Klik area foto atau seret foto ke kotak yang tersedia'])
step_box(doc, 7, 'Isi catatan serah terima (opsional)',
    ['Contoh: "Diserahkan langsung di loket, kondisi barang baik"'])
step_box(doc, 8, 'Klik Konfirmasi Selesai')

note_box(doc,'✅','Serah Terima Tercatat!',
    'Foto dan catatan serah terima tersimpan permanen di sistem.\n'
    'Admin dapat melihat riwayat lengkap di Panel Admin → Riwayat Pengambilan.',
    'F0FDF4',(22,163,74))

doc.add_page_break()

# ═══════════════════════════════════════════════
# 6. CEK STATUS KLAIM
# ═══════════════════════════════════════════════
section_header(doc,'6','Cara Cek Status Klaim')

para(doc,'Untuk memantau status klaim yang sudah diajukan:',11,space_after=8)

step_box(doc, 1, 'Klik menu 🤝 Pengambilan di navbar')
step_box(doc, 2, 'Klik tab Cek Status')
step_box(doc, 3, 'Masukkan nomor klaim, nama, atau nomor kontak pengklaim')
step_box(doc, 4, 'Klik tombol Cek')
step_box(doc, 5, 'Informasi klaim akan muncul:',
    ['Nama pengklaim & barang yang diklaim',
     'Status klaim (Selesai)',
     'Foto bukti pengambil yang diupload saat klaim'])

doc.add_paragraph()

# ═══════════════════════════════════════════════
# 7. LOGOUT
# ═══════════════════════════════════════════════
section_header(doc,'7','Cara Logout')

para(doc,'Selalu logout setelah selesai menggunakan aplikasi:',11,space_after=8)

step_box(doc, 1, 'Lihat nama "Rink Operator 🙋" di pojok kanan atas navbar')
step_box(doc, 2, 'Klik tombol Keluar yang ada di sebelahnya')
step_box(doc, 3, 'Anda akan kembali ke tampilan Guest (tanpa login)')

note_box(doc,'🔒','Penting!',
    'Selalu logout jika menggunakan perangkat bersama untuk menjaga keamanan akun.',
    'FEF2F2',(220,38,38))

doc.add_paragraph()

# ═══════════════════════════════════════════════
# REFERENSI CEPAT
# ═══════════════════════════════════════════════
tbl_ref = doc.add_table(rows=1,cols=1)
cell_ref = tbl_ref.cell(0,0); cell_bg(cell_ref,'1E3A8A')
pr = cell_ref.paragraphs[0]; pr.paragraph_format.space_before=Pt(6); pr.paragraph_format.space_after=Pt(6)
rr = pr.add_run('  📋  REFERENSI CEPAT — RINK OPERATOR')
rr.bold=True; rr.font.size=Pt(13); rr.font.color.rgb=RGBColor(255,255,255)
doc.add_paragraph()

tbl_q = doc.add_table(rows=6, cols=2); tbl_q.style='Table Grid'
cell_bg(tbl_q.rows[0].cells[0],'1E40AF'); cell_bg(tbl_q.rows[0].cells[1],'1E40AF')
tbl_q.rows[0].cells[0].text='  Tugas'; tbl_q.rows[0].cells[1].text='  Langkah Cepat'
for cell in tbl_q.rows[0].cells:
    for p in cell.paragraphs:
        for r in p.runs: r.bold=True; r.font.color.rgb=RGBColor(255,255,255); r.font.size=Pt(10)

quick = [
    ('Lapor barang ditemukan','✅ Barang Ditemukan → + Laporkan Temuan → Isi form → Kirim'),
    ('Proses klaim pemilik','🤝 Pengambilan → Ajukan Klaim → Step 1,2,3 → Kirim Klaim'),
    ('Upload bukti serah terima','🤝 Pengambilan → Cek Status → Cari klaim → Upload Bukti Serah Terima'),
    ('Cek status klaim','🤝 Pengambilan → Cek Status → Masukkan nama/kontak → Cek'),
    ('Lihat detail & kontak barang','Klik kartu barang mana saja → Detail muncul termasuk foto & kontak'),
]
colors = ['DBEAFE','EFF6FF','DBEAFE','EFF6FF','DBEAFE']
for i,(t,s) in enumerate(quick):
    tbl_q.rows[i+1].cells[0].text=f'  {t}'; tbl_q.rows[i+1].cells[1].text=f'  {s}'
    cell_bg(tbl_q.rows[i+1].cells[0],colors[i]); cell_bg(tbl_q.rows[i+1].cells[1],'FFFFFF')
    for cell in tbl_q.rows[i+1].cells:
        for p in cell.paragraphs:
            for r in p.runs: r.font.size=Pt(10)
    tbl_q.rows[i+1].cells[0].paragraphs[0].runs[0].bold=True

doc.add_paragraph()
note_box(doc,'📞','Butuh Bantuan?',
    'Hubungi pengelola sistem BXRink:\n'
    '📱 WhatsApp: +62 812-3456-7890\n'
    '🕙 Jam Operasional: Senin – Minggu, 10.00 – 17.00 WIB\n'
    '🌐 Aplikasi: stirring-yeot-a170188.netlify.app',
    'F0F9FF',(8,145,178))

# Footer
p_ft = doc.add_paragraph('Panduan Rink Operator — Lost & Found BXRink Ice Skating  |  Dokumen ini bersifat internal')
p_ft.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in p_ft.runs: r.font.size=Pt(9); r.font.color.rgb=RGBColor(148,163,184)

out = r'C:\Users\Lenovo\OneDrive\Documents\lostandfound\Panduan_RinkOperator.docx'
doc.save(out)
print('SUCCESS:', out)
